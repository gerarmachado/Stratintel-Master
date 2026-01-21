import streamlit as st
import google.generativeai as genai
from langchain_google_genai import ChatGoogleGenerativeAI
import pypdf
from docx import Document
from fpdf import FPDF
from io import BytesIO
import requests
from bs4 import BeautifulSoup
from youtube_transcript_api import YouTubeTranscriptApi
import yt_dlp
import os
import time
import datetime
from langchain_community.tools import DuckDuckGoSearchRun
import graphviz

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="StratIntel (Master)", page_icon="♟️", layout="wide")

# ==========================================
# 🔐 SISTEMA DE LOGIN
# ==========================================
def check_password():
    def password_entered():
        if st.session_state["username"] in st.secrets["passwords"] and \
           st.session_state["password"] == st.secrets["passwords"][st.session_state["username"]]:
            st.session_state["password_correct"] = True
            del st.session_state["password"]
        else:
            st.session_state["password_correct"] = False

    if st.session_state.get("password_correct", False):
        return True

    st.markdown("## ♟️ StratIntel: Acceso Restringido")
    st.text_input("Usuario", key="username")
    st.text_input("Contraseña", type="password", on_change=password_entered, key="password")
    
    if "password_correct" in st.session_state and not st.session_state["password_correct"]:
        st.error("❌ Credenciales inválidas")
    return False

if not check_password():
    st.stop()

# ==========================================
# ⚙️ CONFIGURACIÓN Y MODELO
# ==========================================
API_KEY_FIJA = "" 
if "GOOGLE_API_KEY" in st.secrets:
    API_KEY_FIJA = st.secrets["GOOGLE_API_KEY"]

MODELO_ACTUAL = "gemini-2.5-flash"  

# ==========================================
# 🧠 BASE DE DATOS MAESTRA (GRAND UNIFIED STRATEGY)
# ==========================================
DB_CONOCIMIENTO = {
    "✨ RECOMENDACIÓN AUTOMÁTICA": {
        "desc": "La IA decide la mejor estrategia basándose en el contenido.",
        "preguntas": ["Identifica los hallazgos estratégicos más críticos.", "Realiza una evaluación integral de riesgos.", "Genera un Resumen Ejecutivo (BLUF).", "¿Cuáles son las anomalías o patrones ocultos más relevantes?"]
    },

    # =========================================================================
    # 🌍 BLOQUE 1: ESTRUCTURA, HISTORIA Y PODER (REFINADO)
    # =========================================================================
    
    # 1.1 META-TEORÍA Y FUERZAS PROFUNDAS -------------------------------------
    "--- 1.1 FUERZAS PROFUNDAS Y TEORÍA ---": { "desc": "", "preguntas": [] },

    "Renouvin & Duroselle (Fuerzas Profundas)": {
        "desc": "Las fuerzas subyacentes (geografía, demografía, economía) vs la decisión del estadista.",
        "preguntas": [
            "Fuerzas Profundas Materiales: ¿Cómo la geografía, demografía o economía limitan inevitablemente la acción política (Determinismo)?",
            "Fuerzas Profundas Espirituales: ¿Qué papel juegan los nacionalismos, el sentimiento colectivo o la psicología de masas en este conflicto?",
            "El Estadista vs la Fuerza: ¿El líder está moldeando la historia o simplemente está siendo arrastrado por corrientes profundas que no controla?"
        ]
    },
    "Dougherty & Pfaltzgraff (Teorías en Pugna)": {
        "desc": "Marco comparativo de teorías (Ambientales, Sistémicas, Decisorias).",
        "preguntas": [
            "Teorías Ambientales: ¿El conflicto es inevitable debido a la escasez de recursos o presiones ecológicas?",
            "Nivel de Análisis: ¿La causa raíz está en el Individuo (líder), el Estado (régimen) o el Sistema (anarquía)?",
            "Integración vs Desintegración: ¿Las fuerzas tecnológicas y económicas están uniendo a los actores o fragmentándolos en bloques hostiles?"
        ]
    },
    
    # 1.2 GRAN ESTRATEGIA Y CICLOS IMPERIALES ---------------------------------
    "Jean-Baptiste Duroselle (Todo Imperio Perecerá)": {
        "desc": "Ciclos vitales de las entidades políticas y sus causas de muerte.",
        "preguntas": [
            "Causa Exógena vs Endógena: ¿La amenaza principal proviene de una invasión externa o de la descomposición interna?",
            "Pérdida de Energía Creadora: ¿La sociedad ha dejado de innovar y se ha vuelto rígida y burocrática?",
            "Expansión Incontrolada: ¿Ha superado el Estado su 'radio de acción eficaz', volviéndose ingobernable?"
        ]
    },
    "Robert Kaplan (La Venganza de la Geografía)": {
        "desc": "El mapa como destino y las restricciones físicas del poder.",
        "preguntas": [
            "El Mapa del Alivio: ¿Cómo las montañas, ríos y desiertos imponen límites físicos insuperables a la ideología política?",
            "Zonas de Choque: ¿Está el conflicto ocurriendo en una 'zona de aplastamiento' (shatterbelt) histórica inevitable?",
            "Imperativos Geográficos: ¿Qué acciones está tomando el Estado simplemente porque su geografía se lo exige (salida al mar, defensa de fronteras llanas)?"
        ]
    },
    "Paul Kennedy (Auge y Caída de las Grandes Potencias)": {
        "desc": "Sobrestiramiento imperial (Imperial Overstretch).",
        "preguntas": [
            "Sobrestiramiento Imperial: ¿Están los compromisos militares y estratégicos del actor superando su capacidad económica para sostenerlos?",
            "Base Económica vs Poder Militar: ¿Se está sacrificando la inversión productiva interna para financiar la seguridad externa?",
            "Erosión Relativa: ¿El declive es absoluto o simplemente los rivales están creciendo más rápido?"
        ]
    },
    
    # 1.3 REALISMO PURO Y GEOPOLÍTICA DURA ------------------------------------
    "Halford Mackinder (Teoría del Heartland)": {
        "desc": "El control de la Isla Mundial y el Pivote Geográfico.",
        "preguntas": [
            "Pivote Geográfico: ¿Quién controla actualmente el 'Heartland' (Eurasia central)?",
            "Cinturón Interior: ¿Hay un conflicto por el control de las zonas costeras que rodean el Heartland?",
            "Contención Terrestre: ¿Se está usando el poder terrestre para negar el acceso a las potencias marítimas?"
        ]
    },
    "Nicholas Spykman (Teoría del Rimland)": {
        "desc": "El control de los bordes anfibios (Rimland).",
        "preguntas": [
            "Anfibia Estratégica: Analiza el conflicto en las zonas costeras/peninsulares (Rimland).",
            "Cerco: ¿Están las potencias tratando de rodear al actor central desde el mar?",
            "Valor de las Alianzas: ¿Qué alianzas en el borde euroasiático son vitales para mantener el equilibrio?"
        ]
    },
    "Hans Morgenthau (Realismo Clásico Integral)": {
        "desc": "Los 6 Principios del Realismo Político y el Interés como Poder.",
        "preguntas": [
            "Leyes Objetivas: ¿Qué fuerzas inherentes a la naturaleza humana (egoísmo, dominio) están impulsando este conflicto?",
            "Interés y Poder: Define el 'Interés Nacional' de los actores en términos de poder, no de moralidad.",
            "Supervivencia del Estado: ¿Está la integridad territorial o política del Estado en riesgo directo?",
            "Autonomía de la Esfera Política: Analiza la decisión desde una lógica puramente política, ignorando consideraciones económicas o legales secundarias."
        ]
    },
    "Kenneth Waltz (Neorrealismo / Imágenes)": {
        "desc": "Las Tres Imágenes (Hombre, Estado, Sistema) y la Estructura Anárquica.",
        "preguntas": [
            "Tercera Imagen (Sistémica): ¿Cómo la anarquía internacional y la distribución de poder (polaridad) obligan al actor a actuar así?", 
            "Polaridad: ¿Cómo afecta la distribución de capacidades (unipolar/multipolar)?",
            "Segunda Imagen (Estatal): ¿Es el régimen político interno irrelevante para la política exterior en este caso?",
            "Equilibrio de Poder: ¿Está el actor haciendo 'Balancing' (aliarse contra el fuerte) o 'Bandwagoning' (unirse al fuerte)?",
            "Principio de Autoayuda: ¿Qué medidas unilaterales está tomando el actor para garantizar su propia seguridad? ¿El comportamiento es defensivo (seguridad) u ofensivo (poder)?"
        ]
    },
    "John Mearsheimer (Realismo Ofensivo)": {
        "desc": "La Tragedia de las Grandes Potencias y la Hegemonía.",
        "preguntas": [
            "Búsqueda de Hegemonía: ¿Está el actor intentando convertirse en el Hegemon regional para asegurar su supervivencia? ¿Está aprovechando oportunidades para alterar el status quo?",
            "Poder Detenedor del Agua: ¿Cómo la geografía (océanos, montañas) limita la proyección de poder del actor? Evalúa el potencial de poder latente (economía/población) vs poder militar actual.",
            "Maximizador de Poder: ¿Está el actor aprovechando cada oportunidad para debilitar a sus rivales potenciales? ¿Cómo está maximizando su poder relativo a expensas de sus vecinos?",
            "Estrategia de 'Buck-Passing': ¿Está intentando que otro estado asuma el costo de contener al agresor?"
        ]
    },
    "Stephen Walt & Robert Jervis (Realismo Defensivo)": {
        "desc": "Equilibrio de Amenazas y Dilema de Seguridad.",
        "preguntas": [
            "Teoría del Equilibrio de Amenazas: Evalúa la amenaza combinando: 1) Poder Agregado, 2) Geografía, 3) Capacidad Ofensiva, 4) Intenciones Agresivas. ¿Quién es percibido como el más amenazante (no solo el más fuerte)?",
            "Dilema de Seguridad: ¿Las medidas defensivas de un actor están siendo malinterpretadas como ofensivas por el otro?",
            "Espiral de Conflicto: ¿Cómo una acción defensiva ha provocado una reacción hostil involuntaria? ¿Las intenciones agresivas son reales o producto de la incertidumbre sistémica?"
        ]
    },
    "Realismo Neoclásico (Schweller)": {
        "desc": "El sistema presiona, pero la política interna decide.",
        "preguntas": [
            "¿Qué variables domésticas están filtrando o bloqueando la respuesta al sistema internacional?",
            "¿Es el estado 'coherente' o están las élites fragmentadas?",
            "¿Tiene el gobierno la capacidad extractiva para movilizar recursos ante la amenaza?"
        ]
    },
    "Realismo Periférico (Carlos Escudé)": {
        "desc": "Estrategia de supervivencia para estados dependientes (Sur Global).",
        "preguntas": [
            "Costo-Beneficio de la Soberanía: ¿El costo de confrontar al Hegemon supera los beneficios para el bienestar ciudadano?",
            "Política de Alineamiento: ¿Debería el estado adoptar un perfil bajo o alinearse para obtener recursos y evitar sanciones?",
            "Evaluación de Autonomía: ¿Se está sacrificando el desarrollo económico por una retórica nacionalista vacía?"
        ]
    },

    # -------------------------------------------------------------------------
    # 🤝 BLOQUE 2: ESCUELA LIBERAL Y CONSTRUCTIVISTA (INSTITUCIONES E IDENTIDAD)
    # -------------------------------------------------------------------------
    "--- LIBERALISMO, IDENTIDAD ---": { "desc": "", "preguntas": [] },

    "Joseph Nye (Poder Multidimensional 3D)": {
        "desc": "Soft Power, Smart Power y el Tablero de Ajedrez Tridimensional.",
        "preguntas": [
            "Dimensión Soft Power: ¿Qué activos de cultura, valores o políticas otorgan atracción y legitimidad al actor?",
            "Dimensión Smart Power: ¿Está combinando eficazmente la coerción (Hard) con la persuasión (Soft)?",
            "Tablero Superior (Militar): Analiza la distribución de poder militar (¿Unipolar?).",
            "Tablero Medio (Económico): Analiza la distribución económica (¿Multipolar?).",
            "Tablero Inferior (Transnacional): ¿Qué actores no estatales (Hackers, ONGs, Terrorismo) actúan fuera del control estatal?"
        ]
    },
    "Immanuel Kant (Triángulo de la Paz Liberal)": {
        "desc": "Paz Democrática, Interdependencia Económica e Instituciones.",
        "preguntas": [
            "Paz Democrática: ¿Son los actores democracias? (Si lo son, la probabilidad de guerra disminuye drásticamente).",
            "Interdependencia Económica: ¿El nivel de comercio mutuo hace que la guerra sea demasiado costosa?",
            "Organizaciones Internacionales: ¿Pertenecen a instituciones comunes que medien el conflicto?",
            "Derecho Cosmopolita: ¿Existe un respeto supranacional por los derechos de los ciudadanos?"
        ]
    },
    "Keohane & Nye (Neoliberalismo Institucional)": {
        "desc": "Interdependencia Compleja y Regímenes Internacionales.",
        "preguntas": [
            "Canales Múltiples: ¿Existen conexiones entre sociedades (no solo entre gobiernos)? ¿Qué instituciones facilitan la cooperación?",
            "Ausencia de Jerarquía: ¿Están los temas militares subordinados a temas económicos o ecológicos en esta crisis?",
            "Interdependencia Compleja: ¿Los vínculos económicos hacen la guerra irracional?",
            "Regímenes Internacionales: ¿Qué normas o reglas implícitas gobiernan las expectativas? ¿Existe un régimen internacional que regule este conflicto?"
        ]
    },
    "Alexander Wendt (Constructivismo Social)": {
        "desc": "La anarquía es lo que los estados hacen de ella.",
        "preguntas": [
            "Culturas de la Anarquía: ¿El sistema es Hobbesiano (Enemigos), Lockeano (Rivales) o Kantiano (Amigos)?",
            "Estructura Ideacional: ¿Cómo las identidades históricas y normas sociales definen el interés nacional?",
            "Ciclo de Refuerzo: ¿Cómo las interacciones pasadas han construido la percepción actual de 'amenaza'?",
            "Normas Internacionales: ¿Qué normas están constriñendo o habilitando la acción?"
        ]
    },
    "Samuel Huntington (Choque de Civilizaciones)": {
        "desc": "Conflictos de identidad cultural y religiosa.",
        "preguntas": [
            "Líneas de Falla: ¿Ocurre el conflicto en la frontera entre dos civilizaciones distintas?",
            "Núcleo Identitario: ¿Es el núcleo del conflicto la identidad religiosa o cultural?",
            "Síndrome del País Pariente (Kin-Country): ¿Están otros estados interviniendo por lealtad cultural/religiosa?",
            "Occidente vs El Resto: ¿Es una reacción contra la imposición de valores occidentales?"
        ]
    },

    # =========================================================================
    # ⚔️ BLOQUE 3: ESTRATEGIA MILITAR Y TRANSFORMACIÓN DE LA GUERRA
    # =========================================================================
    "--- ARTE DE LA GUERRA Y NUEVOS CONFLICTOS ---": { "desc": "", "preguntas": [] },

    "B.H. Liddell Hart (La Estrategia de Aproximación Indirecta)": {
        "desc": "Evitar la fortaleza, atacar la debilidad, dislocar al enemigo.",
        "preguntas": [
            "Línea de Menor Resistencia: ¿Está el actor atacando donde el enemigo menos lo espera (física o psicológicamente)?",
            "Dislocación: ¿Las maniobras han logrado separar al enemigo de su base, suministros o equilibrio mental antes del combate?",
            "Objetivos Alternativos: ¿Tiene el plan flexibilidad para cambiar de objetivo y mantener al enemigo en dilema?"
        ]
    },
    "Martin van Creveld (La Transformación de la Guerra)": {
        "desc": "Guerra No-Trinitaria y conflictos de baja intensidad.",
        "preguntas": [
            "Ruptura de la Trinidad: ¿El conflicto ignora la distinción clásica entre Gobierno, Ejército y Pueblo?",
            "Actores No Estatales: ¿Son las facciones, tribus o señores de la guerra más relevantes que el Estado?",
            "Guerra por la Existencia: ¿Se lucha por intereses políticos racionales o por mera supervivencia e identidad?"
        ]
    },
    "Mary Kaldor (Las Nuevas Guerras)": {
        "desc": "Conflictos post-Guerra Fría: Identidad + Globalización + Criminalidad.",
        "preguntas": [
            "Política de Identidad: ¿Se moviliza a la gente basándose en etiquetas étnicas/religiosas en lugar de ideología?",
            "Métodos de Terror: ¿Es el desplazamiento forzado y el ataque a civiles el objetivo central, no un daño colateral?",
            "Economía Depredadora: ¿Se financia la guerra mediante saqueo, mercado negro o ayuda humanitaria desviada?"
        ]
    },
    "Sun Tzu (El Arte de la Guerra)": {
        "desc": "Engaño, velocidad y victoria sin combate.",
        "preguntas": [
            "El Engaño: ¿Toda la operación se basa en una finta o distracción?",
            "Ganar sin luchar: ¿Está el actor logrando sus objetivos políticos sin uso cinético de fuerza?",
            "Conocimiento: ¿Conoce el actor al enemigo y a sí mismo?", 
            "Terreno: ¿Es el terreno mortal, disperso o clave? ¿Cómo afecta la maniobra?"
        ]
    },
    "Carl von Clausewitz (La Guerra Absoluta)": {
        "desc": "La guerra como continuación de la política.",
        "preguntas": [
            "Trinidad Paradójica: Analiza la relación entre Pasión (Pueblo), Probabilidad (Ejército) y Razón (Gobierno).",
            "Niebla y Fricción: ¿Qué imprevistos están ralentizando la operación?",
            "Centro de Gravedad (COG): ¿Cuál es la fuente de poder del enemigo que, si cae, todo el sistema colapsa?",
            "Política: ¿Es esta acción militar coherente con el objetivo político final?"
        ]
    },
    "Guerra Híbrida (Doctrina Gerasimov)": {
        "desc": "Sincronización de medios militares y no militares.",
        "preguntas": [
            "Fase Latente: ¿Se usa desinformación para desestabilizar antes del conflicto?",
            "Fuerzas Proxy: ¿Se utilizan actores no estatales para negar responsabilidad?",
            "Guerra Económica/Informativa: ¿Es el ataque principal cinético (bombas) o no cinético (sanciones/hackeos)?",
            "Dominio de la Información: ¿Es el ataque informativo más devastador que el físico?"
        ]
    },
    "Qiao Liang & Wang Xiangsui (Guerra Irrestricta)": {
        "desc": "Todo es un arma: leyes, economía, drogas, medios.",
        "preguntas": [
            "Desbordamiento del Campo de Batalla: ¿Se está usando el sistema legal (Lawfare) como arma?",
            "Guerra Financiera: ¿Se están atacando las monedas o mercados del adversario?",
            "Guerra Cultural: ¿Se están atacando los valores fundacionales de la sociedad objetivo?"
        ]
    },

    # =========================================================================
    # 💰 BLOQUE 4: GEOECONOMÍA, TRANSNACIONALISMO Y ANARQUÍA
    # =========================================================================
    "--- ECONOMÍA ILÍCITA Y CAOS ---": { "desc": "", "preguntas": [] },

    "Moisés Naím (Ilícito y el Fin del Poder)": {
        "desc": "El lado oscuro de la globalización y la erosión del Estado.",
        "preguntas": [
            "Las Cinco Guerras: Analiza el tráfico de: 1) Drogas, 2) Armas, 3) Personas, 4) Propiedad Intelectual, 5) Dinero sucio.",
            "Micropoderes: ¿Están actores pequeños y ágiles burlando las defensas de grandes burocracias estatales?",
            "Estado Hueco: ¿Tienen las instituciones la fachada de gobierno pero están carcomidas por redes criminales?"
        ]
    },
    "Robert Kaplan (La Anarquía que Viene)": {
        "desc": "Escasez, tribalismo y erosión de fronteras.",
        "preguntas": [
            "Estrés de Recursos: ¿Es la escasez de agua, tierra o comida el motor oculto del conflicto?",
            "Retribalización: ¿Están colapsando las identidades nacionales en favor de lealtades de clan o secta?",
            "Fronteras Porosas: ¿El mapa político oficial ha dejado de representar la realidad del control territorial?"
        ]
    },
    "Edward Luttwak (Geoeconomía)": {
        "desc": "La lógica del conflicto con la gramática del comercio.",
        "preguntas": [
            "Armamentalización del Comercio: ¿Se usan aranceles o bloqueos como armas?",
            "Predación de Inversiones: ¿Está un estado adquiriendo infraestructura crítica del rival?",
            "Soberanía Tecnológica: ¿Se está bloqueando el acceso a tecnología clave?"
        ]
    },

    # =========================================================================
    # 🤝 BLOQUE 5: NEGOCIACIÓN, JUEGOS Y CONFLICTO
    # =========================================================================
    "--- ESTRATEGIA DE INTERACCIÓN ---": { "desc": "", "preguntas": [] },

    "Thomas Schelling (La Estrategia del Conflicto)": {
        "desc": "Disuasión, Compulsión y la Racionalidad de lo Irracional.",
        "preguntas": [
            "Compulsión vs Disuasión: ¿Se intenta impedir una acción (Disuasión) o forzar a que ocurra (Compulsión)?",
            "Puntos Focales (Schelling Points): ¿Existe una solución obvia donde convergerán las expectativas de ambos sin comunicarse?",
            "La Racionalidad de la Irracionalidad: ¿Se está fingiendo locura o descontrol para obligar al otro a ceder?",
            "Quemar los Barcos: ¿El actor se ha quitado a sí mismo la opción de retroceder para hacer creíble su amenaza?"
        ]
    },
    "William Ury (Cómo Negociar sin Ceder)": {
        "desc": "Negociación basada en principios y superación de bloqueos.",
        "preguntas": [
            "Intereses vs Posiciones: ¿Qué es lo que realmente quieren (Interés) vs lo que dicen que quieren (Posición)?",
            "MAPAN (BATNA): ¿Cuál es la Mejor Alternativa a un Acuerdo Negociado de cada parte? (Quién tiene más poder de retiro).",
            "Separar a la Persona del Problema: ¿Están las emociones o egos bloqueando la solución técnica?",
            "El Puente de Oro: ¿Se le está ofreciendo al adversario una salida digna para que no pierda la cara?"
        ]
    },
    "Robert Axelrod (Complejidad de la Cooperación)": {
        "desc": "Teoría de Juegos, Evolución de la Cooperación y Normas.",
        "preguntas": [
            "El Dilema del Prisionero: ¿Existen incentivos estructurales que hacen racional la traición individual?",
            "Estrategia Tit-for-Tat: ¿Está el actor respondiendo con reciprocidad estricta? ¿Está respondiendo proporcionalmente o escalando?",
            "La Sombra del Futuro: ¿Es la interacción lo suficientemente duradera para fomentar la cooperación? ¿Tienen expectativas de interactuar nuevamente?",
            "Meta-Normas: ¿Existe presión social o sanciones de terceros para castigar a los desertores?",
            "Detección de Trampas: ¿Qué mecanismos de verificación existen para asegurar el cumplimiento?",
            "Estructura de Pagos: ¿Cómo alterar los incentivos para que cooperar sea más rentable que traicionar?"
        ]
    },
    "Teoría de Juegos (John Nash)": {
        "desc": "Equilibrios matemáticos en la toma de decisiones.",
        "preguntas": [
            "Suma Cero vs Suma Variable: ¿Para que uno gane, el otro debe perderlo todo?",
            "Equilibrio de Nash: ¿Cuál es la situación donde nadie tiene incentivos para cambiar su estrategia?",
            "La Gallina (Chicken Game): ¿Quién cederá primero ante la inminencia del choque?"
        ]
    },

    # -------------------------------------------------------------------------
    # 🧠 BLOQUE 6: TOMA DE DECISIONES Y ANÁLISIS ESTRATÉGICO
    # -------------------------------------------------------------------------
    "--- TOMA DE DECISIONES Y SEGURIDAD ---": { "desc": "", "preguntas": [] },

    "Graham Allison (Los 3 Modelos de Decisión)": {
        "desc": "Análisis de la crisis desde múltiples lentes (La Esencia de la Decisión).",
        "preguntas": [
            "Modelo I (Actor Racional): ¿Cuál es la opción lógica que maximiza beneficios y minimiza costos estratégicos?",
            "Modelo II (Proceso Organizacional): ¿Qué procedimientos estándar (SOPs) y rutinas limitan la flexibilidad del gobierno?",
            "Modelo III (Política Burocrática): ¿Qué agencias o individuos internos están luchando por el poder y cómo afecta esto la decisión final?"
        ]
    },
    "Barry Buzan (Seguridad Integral y Securitización)": {
        "desc": "Los 5 Sectores de Seguridad y la Teoría de la Securitización.",
        "preguntas": [
            "Análisis Multisectorial: Evalúa amenazas en los 5 sectores: Militar, Político, Económico, Societal y Ambiental.",
            "Nivel Sistémico: ¿Cómo influye la anarquía internacional o la polaridad en el conflicto?",
            "Nivel Estatal: ¿Qué presiones burocráticas o nacionales limitan al Estado?",
            "Nivel Individual: ¿El perfil psicológico de los líderes altera la toma de decisiones?",
            "Seguridad Societal: ¿Está amenazada la identidad colectiva (religión, etnia, cultura)?",
            "Actor Securitizador: ¿Quién está declarando el asunto como una 'amenaza existencial'?",
            "Objeto Referente: ¿Qué es exactamente lo que se intenta proteger (El Estado, la Nación, la Economía)?",
            "Medidas Extraordinarias: ¿Se está usando la retórica de seguridad para justificar acciones fuera de la política normal?"
        ]
    },
    "John Boyd (Ciclo OODA)": {
        "desc": "Velocidad de decisión en conflicto (Observar, Orientar, Decidir, Actuar).",
        "preguntas": [
            "Velocidad del Ciclo: ¿Quién está completando su ciclo OODA más rápido?",
            "Fase de Orientación: ¿Cómo los sesgos culturales y la herencia genética moldean la percepción del adversario?",
            "Colapso del Adversario: ¿Cómo podemos generar ambigüedad para aislar al enemigo de su entorno?"
        ]
    },
    
    # =========================================================================
    # 🌐 BLOQUE 7: CIBERINTELIGENCIA Y REDES
    # =========================================================================
    "--- CIBERESPACIO E INFO ---": { "desc": "", "preguntas": [] },

    "Cyber Kill Chain (Lockheed Martin)": {
        "desc": "Fases de una intrusión cibernética.",
        "preguntas": [
            "Reconocimiento: ¿Qué datos se están recolectando antes del ataque?",
            "Armamentización: ¿Cómo se creó el malware o el exploit?",
            "Entrega y Explotación: ¿Fue phishing, USB, vulnerabilidad web?",
            "Acciones sobre Objetivos: ¿Se busca robar datos, destruir sistemas o secuestrar (Ransomware)?"
        ]
    },
    "Teoría del Actor-Red (Latour)": {
        "desc": "Humanos y objetos (algoritmos) tienen agencia.",
        "preguntas": [
            "Agencia Tecnológica: ¿Cómo un algoritmo o plataforma está moldeando el conflicto por sí solo?",
            "Cajas Negras: ¿Qué procesos técnicos se están aceptando sin cuestionar su funcionamiento?",
            "Traducción: ¿Cómo se están redefiniendo los intereses a través de la red?"
        ]
    },
    "Modelo Diamante de Intrusión": {
        "desc": "Relación entre Adversario, Infraestructura, Capacidad y Víctima.",
        "preguntas": [
            "Eje Adversario-Víctima: ¿Cuál es la intención sociopolítica detrás del ataque técnico?",
            "Eje Infraestructura-Capacidad: ¿Qué servidores o IPs (Infraestructura) soportan el malware (Capacidad)?",
            "Pivoteo: ¿Podemos usar la infraestructura detectada para encontrar otras víctimas desconocidas?"
        ]
    },

    # =========================================================================
    # 🧠 BLOQUE 8: PSICOLOGÍA Y MENTE DEL ADVERSARIO
    # =========================================================================
    "--- PSICOLOGÍA OPERATIVA ---": { "desc": "", "preguntas": [] },

    "Robert M. Ryder (Conciencia de Dominio / Domain Awareness)": {
        "desc": "Comprensión holística y cognitiva del entorno operativo total.",
        "preguntas": [
            "Ceguera de Dominio: ¿Qué esfera del entorno (marítima, ciber, espacial, humana) estamos ignorando por falta de sensores?",
            "Fusión de Datos: ¿Se están conectando puntos aislados para formar una imagen operativa común (COP)?",
            "Anticipación Cognitiva: ¿Estamos reaccionando a eventos o previendo flujos en el entorno?",
            "Conciencia Cultural: ¿Entendemos el 'terreno humano' tan bien como el terreno físico?"
        ]
    },
    "Perfilado Dark Triad (Tríada Oscura)": {
        "desc": "Psicopatía, Narcisismo y Maquiavelismo en el liderazgo.",
        "preguntas": [
            "Narcisismo: ¿El líder necesita admiración constante y reacciona con ira a la crítica?",
            "Maquiavelismo: ¿Manipula a aliados y enemigos sin remordimiento?",
            "Psicopatía: ¿Muestra falta total de empatía y toma riesgos impulsivos?",
            "Vulnerabilidad del Ego: ¿Cómo se puede explotar su necesidad de validación?"
        ]
    },
    "Código MICE (Motivaciones de Traición)": {
        "desc": "Money, Ideology, Coercion, Ego.",
        "preguntas": [
            "Dinero (Money): ¿Existen crisis financieras personales?",
            "Ideología (Ideology): ¿Cree el sujeto en una causa superior opuesta?",
            "Coerción (Coercion): ¿Existe material de chantaje (Kompromat)?",
            "Ego: ¿Se siente infravalorado o busca venganza?"
        ]
    },
    "Gustave Le Bon (Psicología de Masas)": {
        "desc": "Comportamiento irracional y contagio emocional.",
        "preguntas": [
            "Contagio Mental: ¿Cómo se propaga la emoción irracional?",
            "Líder de Masas: ¿Quién canaliza el odio o la esperanza de la multitud?",
            "Imágenes Simplistas: ¿Qué eslóganes reemplazan el pensamiento lógico?"
        ]
    },

    # =========================================================================
    # 🔮 BLOQUE 9: PROSPECTIVA Y COMPLEJIDAD
    # =========================================================================
    "--- FUTUROS Y SISTEMAS ---": { "desc": "", "preguntas": [] },

    "Análisis Causal por Capas (CLA - Inayatullah)": {
        "desc": "Deconstrucción profunda de la realidad.",
        "preguntas": [
            "La Letanía: ¿Qué dicen los titulares oficiales?",
            "Causas Sistémicas: ¿Qué estructuras generan el problema?",
            "Visión del Mundo: ¿Qué ideologías sostienen el sistema?",
            "Mito y Metáfora: ¿Cuál es la historia inconsciente detrás de todo?"
        ]
    },
    "Nassim Taleb (Cisne Negro & Antifragilidad)": {
        "desc": "Gestión de lo improbable y el caos.",
        "preguntas": [
            "Cisne Negro: Evento de probabilidad baja e impacto infinito.",
            "Rinoceronte Gris: Amenaza obvia ignorada voluntariamente.",
            "Antifragilidad: ¿Qué actor se beneficia del desorden?"
        ]
    },
    "Análisis de Señales Débiles (Weak Signals)": {
        "desc": "Detección temprana de anomalías.",
        "preguntas": [
            "Ruido Marginal: ¿Qué dato 'irrelevante' se repite?",
            "Ceguera Experta: Identifica qué escenarios están siendo descartados por los expertos oficiales por considerarlos 'imposibles' o 'ridículos'.",
            "Patrones de Rareza: ¿Qué evento rompe la continuidad histórica?"
        ]
    },

    # =========================================================================
    # 🛠️ BLOQUE 10: HERRAMIENTAS TÁCTICAS (SATs)
    # =========================================================================
    "--- HERRAMIENTAS ESTRUCTURADAS ---": { "desc": "", "preguntas": [] },

    "Análisis de Hipótesis en Competencia (ACH)": { "desc": "Matriz científica para evitar sesgos.", "preguntas": ["Generación de Hipótesis.", "Matriz de Evidencia.", "Diagnóstico de Consistencia.", "Refutación."] },
    "Análisis de Actores (Stakeholder Mapping)": { "desc": "Mapa de poder e intereses.", "preguntas": ["Matriz Poder/Interés.", "Vetadores.", "Spoilers (Saboteadores)."] },
    "Matriz CARVER (Selección de Objetivos)": { "desc": "Evaluación de blancos.", "preguntas": ["Criticidad.", "Accesibilidad.", "Recuperabilidad.", "Vulnerabilidad.", "Efecto.", "Reconocibilidad."] },
    "Análisis PMESII-PT (Entorno Operativo)": { "desc": "Análisis holístico.", "preguntas": ["Político/Militar.", "Económico/Social.", "Información/Infraestructura.", "Físico/Tiempo."] },
    "Análisis FODA (SWOT) de Inteligencia": { "desc": "Ofensivo/Defensivo.", "preguntas": ["Amenazas Inminentes.", "Oportunidades.", "Vulnerabilidades Internas.", "Fortalezas."] },
    "Técnica de los 5 Porqués": { "desc": "Búsqueda de Causa Raíz.", "preguntas": ["Síntoma.", "¿Por qué? (x5).", "Falla Sistémica."] },
    "Abogado del Diablo": { "desc": "Desafío de asunciones.", "preguntas": ["Desafío Frontal a la tesis principal.", "Defensa de la postura irracional del adversario."] },
    "Richards J. Heuer (Psicología del Análisis de Inteligencia)": { "desc": "Chequeo de sesgos cognitivos del propio analista.", "preguntas": ["Sesgo de Confirmación: ¿Estamos buscando solo información que confirma nuestra hipótesis y descartando la que la contradice?", "Imagen en Espejo: ¿Estamos asumiendo que el adversario piensa y actúa racionalmente como nosotros?", "Anclaje: ¿Estamos demasiado atados a la primera estimación o dato que recibimos al inicio de la crisis?"
         ]
    }
}

# --- GESTIÓN DE ESTADO ---
if 'api_key' not in st.session_state: st.session_state['api_key'] = ""
if 'texto_analisis' not in st.session_state: st.session_state['texto_analisis'] = ""
if 'origen_dato' not in st.session_state: st.session_state['origen_dato'] = "Ninguno"

# --- FUNCIONES DE PROCESAMIENTO ---
def buscar_en_web(query):
    try:
        search = DuckDuckGoSearchRun()
        return search.run(query)
    except Exception as e: return f"Error web: {e}"

def procesar_archivos_pdf(archivos):
    texto_total = ""
    nombres = []
    for archivo in archivos:
        reader = pypdf.PdfReader(archivo)
        texto_pdf = "".join([p.extract_text() for p in reader.pages])
        texto_total += f"\n--- ARCHIVO: {archivo.name} ---\n{texto_pdf}\n"
        nombres.append(archivo.name)
    return texto_total, str(nombres)

def procesar_archivos_docx(archivos):
    texto_total = ""
    nombres = []
    for archivo in archivos:
        doc = Document(archivo)
        texto_doc = "\n".join([para.text for para in doc.paragraphs])
        texto_total += f"\n--- ARCHIVO: {archivo.name} ---\n{texto_doc}\n"
        nombres.append(archivo.name)
    return texto_total, str(nombres)

def obtener_texto_web(url):
    try:
        h = {'User-Agent': 'Mozilla/5.0'}
        r = requests.get(url, headers=h, timeout=15)
        s = BeautifulSoup(r.content, 'html.parser')
        for script in s(["script", "style"]): script.extract()
        return s.get_text(separator='\n')
    except Exception as e: return f"Error: {e}"

def procesar_youtube(url, api_key):
    vid = url.split("v=")[-1].split("&")[0] if "v=" in url else url.split("/")[-1]
    try:
        t = YouTubeTranscriptApi.get_transcript(vid, languages=['es', 'en'])
        return " ".join([i['text'] for i in t]), "Subtítulos"
    except:
        st.info(f"Multimodal (Audio)...")
        opts = {'format': 'bestaudio/best', 'outtmpl': '%(id)s.%(ext)s', 'postprocessors': [{'key': 'FFmpegExtractAudio','preferredcodec': 'mp3'}], 'quiet': True}
        try:
            with yt_dlp.YoutubeDL(opts) as ydl:
                info = ydl.extract_info(url, download=True)
                fname = f"{info['id']}.mp3"
            genai.configure(api_key=api_key)
            myfile = genai.upload_file(fname)
            while myfile.state.name == "PROCESSING": time.sleep(2); myfile = genai.get_file(myfile.name)
            model = genai.GenerativeModel(MODELO_ACTUAL)
            res = model.generate_content([myfile, "Transcribe el audio."])
            if os.path.exists(fname): os.remove(fname)
            myfile.delete()
            return res.text, "Audio IA"
        except Exception as e: return f"Error: {e}", "Error"

def generar_esquema_graphviz(texto_analisis, api_key):
    """Genera código DOT para visualizar relaciones."""
    try:
        genai.configure(api_key=api_key)
        # Usamos el modelo flash por velocidad
        model = genai.GenerativeModel("gemini-1.5-flash") 
        
        prompt = f"""
        ACTÚA COMO: Experto en Visualización de Datos de Inteligencia.
        OBJETIVO: Convertir el siguiente análisis textual en un DIAGRAMA DE RED o MAPA CONCEPTUAL.
        
        INSTRUCCIONES TÉCNICAS:
        1. Analiza el texto e identifica: Actores clave, Acciones/Relaciones y Conceptos.
        2. Genera EXCLUSIVAMENTE el código en lenguaje DOT (Graphviz).
        3. NO escribas explicaciones, ni markdown (como ```dot), solo el código crudo.
        
        ESTILO DEL GRAFO:
        - Tipo: digraph G {{ rankdir=LR; node [shape=box, style=filled, color=lightblue, fontname="Arial"]; edge [fontname="Arial", fontsize=10]; }}
        - Relaciones: "Actor A" -> "Actor B" [label="acción"];
        
        TEXTO BASE:
        {texto_analisis[:15000]}
        """
        
        res = model.generate_content(prompt)
        # Limpieza brutal para asegurar que solo quede el código DOT
        codigo_dot = res.text.replace("```dot", "").replace("```", "").replace("DOT", "").strip()
        
        # Crear objeto Graphviz
        grafico = graphviz.Source(codigo_dot)
        return grafico, None
        
    except Exception as e:
        return None, f"Error visual: {e}"

# --- FUNCIONES DE REPORTE ---
def limpiar_texto(t):
    if not t: return ""
    reps = {"✨": "", "🚀": "", "⚠️": "[!]", "✅": "[OK]", "🛡️": "", "🔒": "", "🎖️": "", "♟️": "", "⚖️": ""}
    for k,v in reps.items(): t = t.replace(k,v)
    return t.encode('latin-1', 'replace').decode('latin-1')

class PDFReport(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, 'StratIntel Report V16', 0, 1, 'C')
        self.ln(5)
    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 7)
        self.cell(0, 10, 'Generado por IA. Uso Confidencial.', 0, 0, 'C')

def crear_pdf(texto, tecnicas, fuente):
    pdf = PDFReport()
    pdf.add_page()
    pdf.set_font("Arial", "B", 10)
    pdf.multi_cell(0, 5, limpiar_texto(f"Fuente: {fuente}\nTécnicas: {tecnicas}"))
    pdf.ln(5)
    pdf.set_font("Arial", "", 10)
    pdf.multi_cell(0, 5, limpiar_texto(texto))
    return pdf.output(dest='S').encode('latin-1', 'replace')

def crear_word(texto, tecnicas, fuente):
    doc = Document()
    doc.add_heading('StratIntel Intelligence Report', 0)
    doc.add_paragraph(f"Fuente: {fuente}").bold = True
    doc.add_paragraph(f"Técnicas: {tecnicas}").bold = True
    for l in texto.split('\n'):
        if l.startswith('#'): doc.add_heading(l.replace('#','').strip(), level=2)
        else: doc.add_paragraph(l)
    
    aviso = doc.add_paragraph()
    aviso.add_run("\n\n------------------\nAVISO: Generado por IA. Verificar datos.").font.size = 8
    b = BytesIO(); doc.save(b); b.seek(0)
    return b

# --- INTERFAZ ---
st.sidebar.title("♟️ StratIntel")
st.sidebar.caption("Master Edition | Ops Mode")
st.sidebar.markdown("---")

if API_KEY_FIJA:
    st.session_state['api_key'] = API_KEY_FIJA
    genai.configure(api_key=API_KEY_FIJA)
    st.sidebar.success(f"✅ Conectado ({MODELO_ACTUAL})")
else:
    if not st.session_state['api_key']:
        k = st.sidebar.text_input("🔑 API KEY:", type="password")
        if k: st.session_state['api_key'] = k; genai.configure(api_key=k); st.rerun()

# SELECTOR MULTI-TECNICA
st.sidebar.subheader("🎯 Misión")
tecnicas_seleccionadas = st.sidebar.multiselect(
    "Técnicas (Máx 3):",
    options=list(DB_CONOCIMIENTO.keys()),
    max_selections=3
)

temp = st.sidebar.slider("Creatividad", 0.0, 1.0, 0.4)
if st.sidebar.button("🔒 Salir"): del st.session_state["password_correct"]; st.rerun()

st.title("♟️ StratIntel | División de Análisis")
st.markdown("**Sistema de Inteligencia Estratégica (DSS)**")

# CARGA
t1, t2, t3, t4, t5 = st.tabs(["📂 PDFs", "📝 DOCXs", "🌐 Web", "📺 YouTube", "✍️ Manual"])
with t1:
    f = st.file_uploader("PDFs", type="pdf", accept_multiple_files=True)
    if f and st.button("Procesar PDF"):
        t, n = procesar_archivos_pdf(f); st.session_state['texto_analisis']=t; st.session_state['origen_dato']=f"PDFs: {n}"; st.success(f"✅ {len(f)}")
with t2:
    f = st.file_uploader("DOCXs", type="docx", accept_multiple_files=True)
    if f and st.button("Procesar DOCX"):
        t, n = procesar_archivos_docx(f); st.session_state['texto_analisis']=t; st.session_state['origen_dato']=f"DOCXs: {n}"; st.success(f"✅ {len(f)}")
with t3:
    u = st.text_input("URL"); 
    if st.button("Web"): st.session_state['texto_analisis']=obtener_texto_web(u); st.session_state['origen_dato']=f"Web: {u}"; st.success("OK")
with t4:
    y = st.text_input("YouTube")
    if st.button("Video"):
        with st.spinner("..."):
            t,m=procesar_youtube(y,st.session_state['api_key'])
            if m!="Error": st.session_state['texto_analisis']=t; st.session_state['origen_dato']=f"YT: {y}"; st.success("OK")
            else: st.error(t)
with t5:
    m = st.text_area("Manual")
    if st.button("Fijar"): st.session_state['texto_analisis']=m; st.session_state['origen_dato']="Manual"; st.success("OK")

st.markdown("---")
if st.session_state['texto_analisis']:
    with st.expander(f"Fuente Activa: {st.session_state['origen_dato']}"): st.write(st.session_state['texto_analisis'][:1000])

# EJECUCIÓN
st.header("Generación de Informe")

if not st.session_state['api_key'] or not st.session_state['texto_analisis']:
    st.warning("⚠️ Carga datos para comenzar.")
else:
    c1, c2 = st.columns([1, 2])
    with c1:
        if not tecnicas_seleccionadas: st.info("👈 Selecciona técnicas.")
        
        # --- SELECTOR DE PROFUNDIDAD CON MODO OPERACIONAL ---
        profundidad = st.radio(
            "Nivel de Profundidad:", 
            ["🔍 Estratégico (Resumen)", "🎯 Táctico (Todas las preguntas)", "⚙️ Operacional (Selección Específica)"],
            help="Estratégico: Visión general. Táctico: Todas las preguntas del marco. Operacional: Selecciona preguntas manualmente."
        )
        
        # --- LÓGICA DE SELECCIÓN MANUAL (OPERACIONAL) ---
        preguntas_manuales = {}
        if "Operacional" in profundidad and tecnicas_seleccionadas:
            st.info("👇 Selecciona los vectores de análisis:")
            for tec in tecnicas_seleccionadas:
                # Obtenemos las preguntas de TU base de datos exacta
                qs = DB_CONOCIMIENTO.get(tec, {}).get("preguntas", [])
                if qs:
                    sel = st.multiselect(f"Preguntas para {tec}:", qs)
                    preguntas_manuales[tec] = sel
                else:
                    st.warning(f"{tec} no tiene preguntas predefinidas.")
        
        usar_internet = st.checkbox("🌐 Búsqueda Web")
        pir = st.text_area("PIR (Opcional):", height=100)

    with c2:
        if st.button("🚀 EJECUTAR MISIÓN", type="primary", use_container_width=True, disabled=len(tecnicas_seleccionadas)==0):
            try:
                genai.configure(api_key=st.session_state['api_key'])
                model = genai.GenerativeModel(MODELO_ACTUAL)
                ctx = st.session_state['texto_analisis']
                
                # BÚSQUEDA WEB
                contexto_web = ""
                if usar_internet:
                    with st.status("🌐 Buscando...", expanded=True) as s:
                        q = f"{pir} {st.session_state['origen_dato']}" if pir else f"Análisis {st.session_state['origen_dato']}"
                        res_web = buscar_en_web(q)
                        contexto_web = f"\nINFO WEB:\n{res_web}\n"
                        s.update(label="✅ Hecho", state="complete", expanded=False)
                
                # BUCLE DE ANÁLISIS
                informe_final = f"# INFORME\nFECHA: {datetime.datetime.now().strftime('%d/%m/%Y')}\nFUENTE: {st.session_state['origen_dato']}\n\n"
                progreso = st.progress(0)
                
                for i, tec in enumerate(tecnicas_seleccionadas):
                    st.caption(f"Analizando: {tec}...")
                    
                    # LÓGICA DE INYECCIÓN DE PREGUNTAS
                    instruccion_preguntas = ""
                    
                    if "Táctico" in profundidad:
                        qs = DB_CONOCIMIENTO.get(tec, {}).get("preguntas", [])
                        if qs:
                            lista = "\n".join([f"- {p}" for p in qs])
                            instruccion_preguntas = f"\n\nOBLIGATORIO: Responde DETALLADAMENTE a TODAS estas preguntas del marco teórico:\n{lista}"
                        else:
                            instruccion_preguntas = "\n\nINSTRUCCIÓN: Realiza un análisis táctico detallado."

                    elif "Operacional" in profundidad:
                        qs_selec = preguntas_manuales.get(tec, [])
                        if qs_selec:
                            lista = "\n".join([f"- {p}" for p in qs_selec])
                            instruccion_preguntas = f"\n\nOBLIGATORIO: Centra el análisis EXCLUSIVAMENTE en responder estas preguntas seleccionadas:\n{lista}"
                        else:
                            instruccion_preguntas = "\n\n(NOTA: El usuario no seleccionó preguntas específicas. Realiza un análisis general de la técnica)."

                    else: # Estratégico
                        instruccion_preguntas = "\n\nINSTRUCCIÓN: Realiza un análisis estratégico general, fluido y ejecutivo (Resumen Global)."

                    prompt = f"""
                    ACTÚA COMO: Analista de Inteligencia Senior y Experto en Relaciones Internacionales.
                    METODOLOGÍA: {tec}
                    PIR (Requerimiento de Inteligencia): {pir}
                    
                    {instruccion_preguntas}
                    
                    CONTEXTO DOCUMENTAL:
                    {ctx}
                    {contexto_web}
                    
                    FORMATO: Académico, riguroso, citar fuentes del texto.
                    """
                    
                    # RETRY LOGIC
                    intentos = 0
                    exito = False
                    while intentos < 3 and not exito:
                        try:
                            res = model.generate_content(prompt, generation_config=genai.types.GenerationConfig(temperature=temp))
                            informe_final += f"\n\n## 📌 {tec}\n{res.text}\n\n---\n"
                            exito = True
                        except Exception as e:
                            if "429" in str(e):
                                st.warning(f"⚠️ Tráfico alto (429). Esperando 30s... (Intento {intentos+1})")
                                time.sleep(30)
                                intentos += 1
                            else:
                                st.error(f"Error: {e}")
                                break

                    progreso.progress((i + 1) / len(tecnicas_seleccionadas))
                    time.sleep(5) 
                
                st.session_state['res'] = informe_final
                st.session_state['tecnicas_usadas'] = ", ".join(tecnicas_seleccionadas)
                st.success("✅ Informe Generado")
                st.markdown("---")
                st.subheader("🎨 Visual de Inteligencia")
                    # --- MÓDULO VISUAL (NUEVO) ---
                st.markdown("---")
                st.subheader("🎨 Esquema")
            
                if st.button("🗺️ Generar Esquema Táctico", type="secondary"):
                    with st.spinner("Diseñando arquitectura del conflicto..."):
                        # Llamamos a la función nueva
                        grafico, error_vis = generar_esquema_graphviz(st.session_state['res'], st.session_state['api_key'])
                    
                        if grafico:
                            # 1. Mostrar en pantalla
                            st.graphviz_chart(grafico)
                        
                            # 2. Botón de Descarga (PNG)
                            try:
                                # Renderizamos el gráfico a bytes en formato PNG
                                png_bytes = grafico.pipe(format='png')
                            
                                st.download_button(
                                    label="💾 Descargar Mapa (Imagen PNG)",
                                    data=png_bytes,
                                    file_name="Mapa_StratIntel.png",
                                    mime="image/png"
                                )
                            except Exception as e:
                                st.warning(f"Se visualiza pero no se puede descargar (Falta binario Graphviz en sistema): {e}")
                        else:
                            st.error(f"No se pudo generar el mapa: {error_vis}")
                st.markdown(informe_final)

            except Exception as e: st.error(f"Error: {e}")

if 'res' in st.session_state:
    st.markdown("---")
    c1, c2 = st.columns(2)
    c1.download_button("Descargar Word", crear_word(st.session_state['res'], st.session_state['tecnicas_usadas'], st.session_state['origen_dato']), "Reporte.docx")
    try: c2.download_button("Descargar PDF", bytes(crear_pdf(st.session_state['res'], st.session_state['tecnicas_usadas'], st.session_state['origen_dato'])), "Reporte.pdf")
    except: pass
    # ==========================================
    # 🎨 MÓDULO DE VISUALIZACIÓN (PERSISTENTE)
    # ==========================================
    st.markdown("---")
    st.header("🎨 Inteligencia Visual")
    
    # Inicializar estado del gráfico si no existe para que no se borre al recargar
    if 'grafico_dot' not in st.session_state:
        st.session_state['grafico_dot'] = None
    
    c_vis1, c_vis2 = st.columns(2)
    
    with c_vis1:
        # Botón para GENERAR (Solo procesa)
        if st.button("🗺️ Generar Mapa de Actores (Esquema)", type="secondary"):
            with st.spinner("Diseñando arquitectura del conflicto..."):
                grafico, error_vis = generar_esquema_graphviz(st.session_state['res'], st.session_state['api_key'])
                
                if grafico:
                    st.session_state['grafico_dot'] = grafico # Guardar en memoria
                else:
                    st.error(f"Error visual: {error_vis}")

    # Renderizado FUERA del botón (Para que persista en pantalla)
    if st.session_state['grafico_dot']:
        st.graphviz_chart(st.session_state['grafico_dot'])
        
        try:
            img_bytes = st.session_state['grafico_dot'].pipe(format='png')
            st.download_button(
                label="💾 Descargar Esquema (PNG)",
                data=img_bytes,
                file_name="Mapa_StratIntel.png",
                mime="image/png"
            )
        except Exception as e:
            st.warning("⚠️ Visualización activa. Para descargar, instala 'graphviz' en packages.txt")












